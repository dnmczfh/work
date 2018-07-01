import requests
from bs4 import BeautifulSoup
from datetime import datetime
import pandas
import os
import shutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def getNewsDetail(newsUrl):
    result = {}
    res = requests.get(newsUrl)
    res.encoding='utf-8'
    soup = BeautifulSoup(res.text,'lxml')
    try:
        result['title'] = soup.select('.main-title')[0].text
        date = soup.select('.date')[0].text.strip()
        result['dt'] = datetime.strptime(date,'%Y年%m月%d日 %H:%M')
        result['source'] = soup.select('.source')[0].text
        #result['sourceUrl'] = soup.select('.source')[0]['href']
        result['keywords'] = ','.join(keyword.text for keyword in soup.select('#keywords a'))
        article = []
        for p in soup.select('#article p'):
            if '相关新闻' in p.text:
                break
            article.append(p.text.strip())
        result['article'] = '\n'.join(article)
    except:
       print('error')
    return result

def clean_chinese_character(text):
    '''处理特殊的中文符号,将其全部替换为'-' 否则在保存时Windows无法将有的中文符号作为路径'''
    chars = chars = ["/", "\"", "'", "·", "。","？", "！", "，", "、", "；", "：", "‘", "’", "“", "”", "（", "）", "…", "–", "．", "《", "》"];
    new_text = ""
    for i in range(len(text)):
        if text[i] not in chars:
            new_text += text[i]
        else:
            new_text += "_"
    return new_text

def create_docx(news_type, title, content):
    '''这里使用python-docx库将新闻的内容生成word文件'''
    document = Document()
    paragraph = document.add_paragraph(title)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.bold = True

    paragraph = document.add_paragraph(content)

    style = paragraph.style
    font = style.font
    font.size = Pt(10)
    #font.name = '华文彩云'

    name = news_type + "-" + clean_chinese_character(title) + ".docx"
    document.save(news_type + "/" + name)

res = requests.get('http://news.sina.com.cn/china')
res.encoding='utf-8'
soup = BeautifulSoup(res.text,'lxml')

news_total = []
for news in soup.select('.news-item'):
    if len(news.select('h2')) > 0:
        h2 = news.select('h2')[0].text
        time = news.select('.time')[0].text
        newsUrl = news.select('a')[0]['href']
        news_detail = getNewsDetail(newsUrl)
        if len(news_detail)>0 :
            news_total.append(news_detail)
        print(h2)

        if len(news_total)>3 :
            break


#删除旧目录
path ='新浪新闻'
print("deleting old dir")
if os.path.exists(path):
    shutil.rmtree(path)
#创建新目录
print("creating dir: ", path)
os.mkdir(path)

# 存储到EXCEL文件中
df = pandas.DataFrame(news_total)
df.to_excel(path + "/" +'sina_news.xlsx')

#把单个新闻存成word文件
for news in news_total:
    create_docx(path,news['title'], news['article'])

# print(df)



