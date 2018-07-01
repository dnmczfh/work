#编写word文件
#需安装python-docx
#python-docx帮助文件官网：https://python-docx.readthedocs.io/en/latest/index.html

from docx import Document
from docx.shared import Inches,Pt,RGBColor
from docx.enum.style import WD_STYLE_TYPE

document = Document()

# styles=document.styles
# print('段落样式：')
# print('\n'.join([s.name for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]))
# print('\n')
#
# print('表格样式：')
# table_styles=[s for s in styles if s.type == WD_STYLE_TYPE.TABLE]
# for style in table_styles:
#     print(style.name)

document.add_heading('文章标题', 0)

p = document.add_paragraph('一个段落是由多个 ')
p.add_run('run').bold = True
p.add_run('联接')
p.add_run('组成的').italic = True

p = document.add_paragraph()
text = '一个人的命运当然要靠自我奋斗，但是也要考虑到历史的进程'
for i, ch in enumerate(text):
    run = p.add_run(ch)
    font = run.font
    font.name = '华文彩云'
    font.size = Pt(20)
    font.color.rgb = RGBColor(i*10%200+55, i*20%200+55, i*30%200+55)


document.add_heading('一级标题', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph('first item in unordered list', style='ListBullet')
document.add_paragraph('first item in unordered list', style='ListBullet')
document.add_paragraph('first item in ordered list', style='ListNumber')
document.add_paragraph('first item in ordered list', style='ListNumber')

styles=document.styles
document.add_heading('段落样式：', level=1)
# paragraph_styles=[s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
# for s in paragraph_styles:
#     document.add_paragraph(s.name, style='ListBullet')
document.add_heading('表格样式：', level=1)
# table_styles=[s for s in styles if s.type == WD_STYLE_TYPE.TABLE]
# for s in table_styles:
#     document.add_paragraph(s.name, style='ListBullet')

document.add_picture('QQ截图20180414161355.png', width=Inches(2.5))
document.add_page_break()

# get table data -------------
#items = get_things_from_database_or_something()
items=[(1,2,3),(4,5,6),(7,8,9)]
table = document.add_table(rows=1, cols=3,style='Colorful List Accent 2')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in items:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item[0])
    row_cells[1].text = str(item[1])
    row_cells[2].text = str(item[2])

recordset=[{'Qty':3,'Name':'Fish','Desc':'Tom'},
           {'Qty':8, 'Name': 'Cheese', 'Desc': 'Jerry'},
           {'Qty':5, 'Name': 'Bacon', 'Desc': 'Garfield'}]

for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item['Qty'])
    row_cells[1].text = str(item['Name'])
    row_cells[2].text = str(item['Desc'])

document.save('demo.docx')

#修改文件中表格中的内容
document=Document('demo.docx')
table=document.tables[0]
for row,obj_row in enumerate(table.rows):
    for col,cell in enumerate(obj_row.cells):
        cell.text='{},{}'.format(row,col)
document.save('demo1.docx')